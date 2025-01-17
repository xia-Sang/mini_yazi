from pathlib import Path
from typing import Optional
import mimetypes  
from PIL import Image  # 处理图片
import docx  # 处理 Word 文档
from PyPDF2 import PdfReader  # 处理 PDF
import chardet  # 处理文本编码

class PreviewHandler:
    """文件预览处理器"""
    
    def __init__(self):
        mimetypes.init()
        
    def get_preview(self, file_path: Path, max_lines: int = 100) -> tuple[str, str]:
        """获取文件预览内容
        返回: (文件类型描述, 预览内容)
        """
        try:
            mime_type, _ = mimetypes.guess_type(str(file_path))
            if mime_type is None:
                return ('未知文件类型', '无法预览此类型的文件')
            
            # 根据MIME类型选择相应的处理方法
            if mime_type.startswith('text/'):
                return self._handle_text(file_path, max_lines)
            elif mime_type.startswith('image/'):
                return self._handle_image(file_path)
            elif mime_type == 'application/pdf':
                return self._handle_pdf(file_path, max_lines)
            elif mime_type in ['application/msword', 
                             'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
                return self._handle_word(file_path, max_lines)
            else:
                return ('二进制文件', f'文件类型: {mime_type}\n不支持预览')
                
        except Exception as e:
            return ('错误', f'预览失败: {str(e)}')
            
    def _handle_text(self, file_path: Path, max_lines: int) -> tuple[str, str]:
        """处理文本文件"""
        # 检测文件编码
        with open(file_path, 'rb') as f:
            raw = f.read()
            result = chardet.detect(raw)
            encoding = result['encoding'] or 'utf-8'
            
        # 读取文件内容
        with open(file_path, 'r', encoding=encoding) as f:
            lines = f.readlines()[:max_lines]
            content = ''.join(lines)
            if len(lines) >= max_lines:
                content += f'\n... (还有更多内容)'
                
        return ('文本文件', content)
        
    def _handle_image(self, file_path: Path) -> tuple[str, str]:
        """处理图片文件"""
        try:
            with Image.open(file_path) as img:
                info = [
                    f'格式: {img.format}',
                    f'大小: {img.size[0]}x{img.size[1]}',
                    f'色彩模式: {img.mode}',
                ]
                if 'dpi' in img.info:
                    info.append(f'DPI: {img.info["dpi"]}')
                    
            return ('图片文件', '\n'.join(info))
        except Exception as e:
            return ('图片文件', f'无法读取图片信息: {str(e)}')
            
    def _handle_pdf(self, file_path: Path, max_lines: int) -> tuple[str, str]:
        """处理PDF文件"""
        try:
            with open(file_path, 'rb') as f:
                pdf = PdfReader(f)
                info = [
                    f'页数: {len(pdf.pages)}',
                    f'标题: {pdf.metadata.get("/Title", "未知")}',
                    f'作者: {pdf.metadata.get("/Author", "未知")}',
                    '\n预览内容:',
                ]
                
                # 获取第一页的文本
                if len(pdf.pages) > 0:
                    text = pdf.pages[0].extract_text()
                    lines = text.split('\n')[:max_lines]
                    info.extend(lines)
                    if len(lines) >= max_lines:
                        info.append('... (还有更多内容)')
                        
            return ('PDF文件', '\n'.join(info))
        except Exception as e:
            return ('PDF文件', f'无法读取PDF: {str(e)}')
            
    def _handle_word(self, file_path: Path, max_lines: int) -> tuple[str, str]:
        """处理Word文档"""
        try:
            doc = docx.Document(file_path)
            paragraphs = []
            for i, para in enumerate(doc.paragraphs):
                if i >= max_lines:
                    paragraphs.append('... (还有更多内容)')
                    break
                if para.text.strip():
                    paragraphs.append(para.text)
                    
            info = [
                f'段落数: {len(doc.paragraphs)}',
                f'页数: {len(doc.sections)}',
                '\n预览内容:',
            ]
            info.extend(paragraphs)
            
            return ('Word文档', '\n'.join(info))
        except Exception as e:
            return ('Word文档', f'无法读取Word文档: {str(e)}') 